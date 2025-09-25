from __future__ import annotations

import os
import re
from dataclasses import dataclass
from typing import Any, Dict, Iterable, List, Tuple

import pdfplumber
from docx import Document

SENT_SPLIT = re.compile(r"(?<=[\.?\!])\s+(?=[A-Z0-9])")


@dataclass
class Chunk:
    doc_id: str
    text: str
    meta: Dict[str, Any]


def _clean_text(value: str) -> str:
    value = value.replace("\u00ad", "")
    value = re.sub(r"\s+", " ", value).strip()
    return value


def read_pdf(path: str) -> List[Tuple[int, str]]:
    pages: List[Tuple[int, str]] = []
    with pdfplumber.open(path) as pdf:
        for idx, page in enumerate(pdf.pages, start=1):
            content = page.extract_text() or ""
            cleaned = _clean_text(content)
            if cleaned:
                pages.append((idx, cleaned))
    return pages


def read_docx(path: str) -> List[Tuple[int, str]]:
    document = Document(path)
    blocks: List[Tuple[int, str]] = []
    buffer: List[str] = []
    for paragraph in document.paragraphs:
        text = _clean_text(paragraph.text)
        if text:
            buffer.append(text)
        elif buffer:
            joined = " ".join(buffer)
            blocks.append((len(blocks) + 1, joined))
            buffer = []
    if buffer:
        joined = " ".join(buffer)
        blocks.append((len(blocks) + 1, joined))
    return blocks


def sentences(text: str) -> List[str]:
    return [sentence.strip() for sentence in SENT_SPLIT.split(text) if sentence.strip()]


def chunk_text(text: str, max_chars: int = 1200, overlap: int = 150) -> List[str]:
    chunks: List[str] = []
    current: List[str] = []
    size = 0
    for sentence in sentences(text):
        if size + len(sentence) > max_chars and current:
            chunk = " ".join(current)
            chunks.append(chunk)
            trailing = " ".join(chunk.split()[-overlap:]) if overlap else ""
            current = [trailing, sentence] if trailing else [sentence]
            size = len(" ".join(current))
        else:
            current.append(sentence)
            size += len(sentence)
    if current:
        chunks.append(" ".join(current))
    return chunks


def _chunk_with_metadata(
    doc_id: str,
    source_label: str,
    origin: Iterable[Tuple[int, str]],
    origin_key: str,
    max_chars: int = 1200,
) -> List[Chunk]:
    chunks: List[Chunk] = []
    for origin_index, text in origin:
        pieces = chunk_text(text, max_chars=max_chars)
        for part_index, piece in enumerate(pieces, start=1):
            meta = {
                "source": source_label,
                origin_key: origin_index,
            }
            if len(pieces) > 1:
                meta["section_chunk"] = part_index
            chunks.append(Chunk(doc_id=doc_id, text=piece, meta=meta))
    return chunks


def build_chunks(path: str, doc_id: str) -> List[Chunk]:
    source_label = doc_id or os.path.basename(path)
    lower_path = path.lower()
    if lower_path.endswith(".pdf"):
        origin = read_pdf(path)
        return _chunk_with_metadata(doc_id, source_label, origin, "page")
    if lower_path.endswith(".docx"):
        origin = read_docx(path)
        if not origin:
            # Fallback: treat the whole document as one block if headings were empty.
            full_document = Document(path)
            cleaned_parts: List[str] = []
            for paragraph in full_document.paragraphs:
                cleaned = _clean_text(paragraph.text)
                if cleaned:
                    cleaned_parts.append(cleaned)
            combined = " ".join(cleaned_parts)
            origin = [(1, combined)] if combined else []
        return _chunk_with_metadata(doc_id, source_label, origin, "section")
    raise ValueError("Unsupported file type. Gunakan berkas .pdf atau .docx")